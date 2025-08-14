from flask import Flask, request, jsonify, session
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
from datetime import datetime, timedelta
import logging
import uuid
import jwt
from functools import wraps

import pytesseract
from PIL import Image
import pdfplumber
import docx
import openpyxl

# Tesseract yolu (Windows)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Database ve config import
from models.database import db, Company, User, Document, DocumentContent, ChatSession, ChatMessage
from config import config

# Flask uygulamasını başlat
app = Flask(__name__)

# Konfigürasyonu yükle
env = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[env])

# Session (sadece gerektiğinde kullanılır; JWT esas)
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False  # Development için
app.config['SESSION_COOKIE_HTTPONLY'] = True

# CORS
CORS(
    app,
    origins=['http://localhost:8000', 'http://127.0.0.1:8000'],
    supports_credentials=True,
    allow_headers=['Content-Type', 'Authorization'],
    methods=['GET', 'POST', 'PUT', 'DELETE', 'OPTIONS'],
)

# DB
db.init_app(app)

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Upload
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx', 'xls', 'xlsx'}
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# TEXT EXTRACTION HELPERS
def _extract_text_pdf(path):
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                text_parts.append(t)
    return "\n".join(text_parts).strip()

def _extract_text_docx(path):
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs).strip()

def _extract_text_xlsx(path):
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                parts.append(" ".join(vals))
    return "\n".join(parts).strip()

def _extract_text_image(path):
    # Basit OCR (dil: Türkçe model yüklüyse 'tur' da deneyebilirsin)
    try:
        return pytesseract.image_to_string(Image.open(path), lang="tur+eng").strip()
    except Exception:
        # Dil modeli yoksa İngilizce düş
        return pytesseract.image_to_string(Image.open(path)).strip()

def _extract_text_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()

def extract_text_by_ext(path, ext):
    ext = (ext or "").lower()
    if ext in ("pdf",):
        return _extract_text_pdf(path)
    if ext in ("docx",):
        return _extract_text_docx(path)
    if ext in ("xlsx", "xls"):
        return _extract_text_xlsx(path)
    if ext in ("png", "jpg", "jpeg", "gif"):
        return _extract_text_image(path)
    if ext in ("txt",):
        return _extract_text_txt(path)
    # desteklenmeyen formatlarda boş döndür
    return ""

# -------- JWT yardımcıları --------
def generate_token(user_id: int, company_id: int) -> str:
    payload = {'user_id': user_id, 'company_id': company_id, 'exp': datetime.utcnow() + timedelta(hours=24)}
    return jwt.encode(payload, app.config['SECRET_KEY'], algorithm='HS256')

def verify_token(token: str):
    try:
        return jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
    except Exception:
        return None

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        raw = request.headers.get('Authorization', '')
        # Debug istersen aç: app.logger.info(f"[AUTH] Raw Authorization: {raw}")
        if raw.startswith('Bearer '):
            raw = raw[7:]
        if not raw:
            return jsonify({'error': 'Token gerekli'}), 401
        payload = verify_token(raw)
        # Debug: app.logger.info(f"[AUTH] Payload: {payload}")
        if not payload:
            return jsonify({'error': 'Geçersiz veya süresi dolmuş token'}), 401
        request.user_id = payload.get('user_id')
        request.company_id = payload.get('company_id')
        if not request.user_id or not request.company_id:
            return jsonify({'error': 'Token içeriği eksik'}), 401
        return f(*args, **kwargs)
    return decorated
# ----------------------------------

@app.route('/')
def home():
    return jsonify({'message': 'Döküman Yönetim Sistemi API', 'version': '1.0.0', 'status': 'active'})

@app.route('/health')
def health_check():
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})

# -------- AUTH --------
@app.route('/api/auth/login', methods=['POST'])
def login():
    try:
        data = request.get_json() or {}
        email = data.get('email')
        password = data.get('password')
        if not email or not password:
            return jsonify({'error': 'Email ve şifre gerekli'}), 400

        user = User.query.filter_by(email=email).first()
        if not user or not user.check_password(password):
            return jsonify({'error': 'Geçersiz email veya şifre'}), 401
        if not user.is_active:
            return jsonify({'error': 'Hesap aktif değil'}), 401

        token = generate_token(user.id, user.company_id)
        # PyJWT bazı sürümlerde bytes döndürebilir; string'e zorlayalım
        if isinstance(token, bytes):
            token = token.decode('utf-8')

        user.last_login = datetime.utcnow()
        db.session.commit()

        return jsonify({
            'message': 'Giriş başarılı',
            'token': token,
            'user': user.to_dict(),
            'company': user.company.to_dict()
        })
    except Exception:
        logger.exception("Login error")
        return jsonify({'error': 'Giriş yapılırken hata oluştu'}), 500

@app.route('/api/auth/register', methods=['POST'])
def register():
    try:
        data = request.get_json() or {}
        required = ['companyName', 'companyEmail', 'username', 'userEmail', 'password']
        for f in required:
            if not data.get(f):
                return jsonify({'error': f'{f} gerekli'}), 400

        if Company.query.filter_by(email=data['companyEmail']).first():
            return jsonify({'error': 'Bu şirket emaili zaten kayıtlı'}), 400

        company = Company(name=data['companyName'], email=data['companyEmail'])
        db.session.add(company); db.session.flush()

        user = User(username=data['username'], email=data['userEmail'], company_id=company.id, role='admin')
        user.set_password(data['password'])
        db.session.add(user); db.session.commit()

        return jsonify({'message': 'Kayıt başarılı', 'company': company.to_dict(), 'user': user.to_dict()}), 201
    except Exception:
        db.session.rollback()
        logger.exception("Register error")
        return jsonify({'error': 'Kayıt olurken hata oluştu'}), 500

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({'message': 'Çıkış başarılı'})

@app.route('/api/auth/me')
@token_required
def get_current_user():
    user = User.query.get(request.user_id)
    if not user:
        return jsonify({'error': 'Kullanıcı bulunamadı'}), 404
    return jsonify({'user': user.to_dict(), 'company': user.company.to_dict()})

# -------- DOCUMENTS --------
@app.route('/api/documents/upload', methods=['POST'])
@token_required
def upload_documents():
    """Dosya yükle + içerik çıkar (OCR/Parser) + DB'ye yaz"""
    try:
        user_id = request.user_id
        company_id = request.company_id

        if 'files' not in request.files:
            return jsonify({'error': 'Dosya seçilmemiş'}), 400

        files = request.files.getlist('files')
        if not files:
            return jsonify({'error': 'Dosya bulunamadı'}), 400

        uploaded = []

        for f in files:
            if not f or f.filename == '':
                continue
            if not allowed_file(f.filename):
                continue

            # 1) Kaydet
            filename = secure_filename(f.filename)
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            unique = f"{uuid.uuid4()}_{filename}"
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], unique)
            f.save(save_path)

            # 2) Document kaydı
            doc = Document(
                filename=unique,
                original_filename=filename,
                file_path=save_path,
                file_type=ext,
                file_size=os.path.getsize(save_path),
                company_id=company_id,
                uploaded_by=user_id
            )
            db.session.add(doc)
            db.session.flush()  # doc.id almak için

            # 3) İçerik çıkar (PDF/DOCX/XLSX/IMG/TXT)
            content_text = ""
            try:
                content_text = extract_text_by_ext(save_path, ext)
            except Exception as e:
                logger.warning(f"Text extraction failed for doc {doc.id}: {e}")

            # 4) İçeriği DocumentContent'e yaz (model alanlarına göre uyarladık)
            try:
                if content_text:
                    dc = DocumentContent(
                        document_id=doc.id,
                        content=content_text
                        # varsa ekstra alanlar: lang, extracted_at vs.
                    )
                    db.session.add(dc)
                    # Document üstündeki durum alanını güncelle (varsa)
                    if hasattr(doc, "is_processed"):
                        doc.is_processed = True
                else:
                    if hasattr(doc, "is_processed"):
                        doc.is_processed = False
            except Exception as e:
                logger.warning(f"Saving content failed for doc {doc.id}: {e}")
                if hasattr(doc, "is_processed"):
                    doc.is_processed = False

            uploaded.append(doc.to_dict())

        db.session.commit()
        return jsonify({'message': f'{len(uploaded)} dosya başarıyla yüklendi', 'files': uploaded}), 201

    except Exception as e:
        db.session.rollback()
        logger.exception("Upload error")
        return jsonify({'error': 'Dosya yüklenirken hata oluştu'}), 500


@app.route('/api/documents', methods=['GET'])
@token_required
def get_documents():
    try:
        company_id = request.company_id
        docs = Document.query.filter_by(company_id=company_id).order_by(Document.created_at.desc()).all()
        return jsonify({'documents': [d.to_dict() for d in docs]})
    except Exception:
        logger.exception("Get documents error")
        return jsonify({'error': 'Belgeler alınırken hata oluştu'}), 500

@app.route('/api/documents/<int:doc_id>', methods=['DELETE'])
@token_required
def delete_document(doc_id: int):
    """JWT ile korunan belge silme"""
    try:
        company_id = request.company_id

        document = Document.query.filter_by(id=doc_id, company_id=company_id).first()
        if not document:
            return jsonify({'error': 'Belge bulunamadı'}), 404

        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception as e:
            logger.warning(f"File deletion warning: {e}")

        db.session.delete(document)
        db.session.commit()
        return jsonify({'message': 'Belge başarıyla silindi'})
    except Exception:
        db.session.rollback()
        logger.exception("Delete document error")
        return jsonify({'error': 'Belge silinirken hata oluştu'}), 500

@app.route('/api/documents/stats')
@token_required
def get_document_stats():
    try:
        company_id = request.company_id
        total = Document.query.filter_by(company_id=company_id).count()
        processed = Document.query.filter_by(company_id=company_id, is_processed=True).count()
        return jsonify({
            'total_documents': total,
            'processed_documents': processed,
            'pending_documents': total - processed
        })
    except Exception:
        logger.exception("Stats error")
        return jsonify({'error': 'İstatistikler alınırken hata oluştu'}), 500

# DB tabloları ve demo veri
def create_tables():
    with app.app_context():
        db.create_all()
        if not Company.query.first():
            demo_company = Company(name='Demo Şirket', email='demo@company.com')
            db.session.add(demo_company); db.session.commit()
            demo_user = User(username='admin', email='admin@demo.com', company_id=demo_company.id, role='admin')
            demo_user.set_password('123456')
            db.session.add(demo_user); db.session.commit()
            logger.info("Demo veri eklendi: admin@demo.com / 123456")

if __name__ == '__main__':
    create_tables()
    app.run(debug=True, host='0.0.0.0', port=5000)
